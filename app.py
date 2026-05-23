import os
import uuid
import json
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import pdfplumber
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from docx import Document
from dotenv import load_dotenv

from gemini_service import extract_text_from_resume, analyze_resume

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'change-me')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['EDITED_FOLDER'] = 'edited'
app.config['ANALYSIS_FOLDER'] = 'analysis'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

for folder in ['UPLOAD_FOLDER', 'EDITED_FOLDER', 'ANALYSIS_FOLDER']:
    os.makedirs(app.config[folder], exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============ HOME = ANALYZER ============
@app.route('/')
def index():
    return render_template('index.html')


# ============ ANALYZE RESUME ============
@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    role = request.form.get('role', '').strip()
    job_description = request.form.get('job_description', '').strip()

    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Use PDF or DOCX'}), 400
    if not role or not job_description:
        return jsonify({'error': 'Role and Job Description required'}), 400

    filename = secure_filename(file.filename)
    unique_id = str(uuid.uuid4())
    ext = filename.rsplit('.', 1)[1].lower()
    saved_name = f"{unique_id}.{ext}"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], saved_name)
    file.save(filepath)

    try:
        resume_text = extract_text_from_resume(filepath, ext)
        if not resume_text:
            return jsonify({'error': 'Could not extract text from resume'}), 400

        analysis = analyze_resume(resume_text, role, job_description)

        if 'error' in analysis:
            return jsonify(analysis), 500

        analysis_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{unique_id}.json")
        with open(analysis_path, 'w', encoding='utf-8') as f:
            json.dump({
                'role': role,
                'job_description': job_description,
                'analysis': analysis
            }, f, indent=2)

        return jsonify({
            'success': True,
            'file_id': unique_id,
            'ext': ext,
            'editor_url': f'/editor/{unique_id}/{ext}',
            'ats_score': analysis.get('ats_score', 0)
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============ EDITOR PAGE ============
@app.route('/editor/<file_id>/<ext>')
def editor(file_id, ext):
    analysis_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{file_id}.json")
    has_analysis_js = 'true' if os.path.exists(analysis_path) else 'false'
    return render_template(
        'editor.html',
        file_id=file_id,
        ext=ext,
        has_analysis_js=has_analysis_js
    )


# ============ GET ANALYSIS DATA ============
@app.route('/get_analysis/<file_id>')
def get_analysis(file_id):
    analysis_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{file_id}.json")
    if not os.path.exists(analysis_path):
        return jsonify({'error': 'No analysis found'}), 404
    with open(analysis_path, 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))


# ============ EXTRACT FILE CONTENT ============
@app.route('/extract/<file_id>/<ext>')
def extract_content(file_id, ext):
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.{ext}")
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    if ext == 'pdf':
        pages = []
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_data = {
                    'page': page_num,
                    'width': float(page.width),
                    'height': float(page.height),
                    'texts': []
                }
                try:
                    words = page.extract_words(extra_attrs=['size', 'fontname'])
                    for word in words:
                        page_data['texts'].append({
                            'text': word['text'],
                            'x': float(word['x0']),
                            'y': float(word['top']),
                            'x2': float(word['x1']),
                            'y2': float(word['bottom']),
                            'font_size': float(word.get('size', 11)),
                            'font': word.get('fontname', 'Helvetica')
                        })
                except Exception as e:
                    print(f"Error extracting words: {e}")
                pages.append(page_data)
        return jsonify({'type': 'pdf', 'pages': pages})

    elif ext in ('docx', 'doc'):
        doc = Document(filepath)
        paragraphs = [{
            'index': i,
            'text': p.text,
            'style': p.style.name if p.style else 'Normal'
        } for i, p in enumerate(doc.paragraphs)]
        return jsonify({'type': 'docx', 'paragraphs': paragraphs})

    return jsonify({'error': 'Unsupported'}), 400


# ============ PDF PREVIEW IMAGE ============
@app.route('/preview/<file_id>/<ext>/<int:page>')
def preview_page(file_id, ext, page):
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.{ext}")
    if ext == 'pdf':
        try:
            with pdfplumber.open(filepath) as pdf:
                pil_image = pdf.pages[page].to_image(resolution=120).original
                buf = BytesIO()
                pil_image.save(buf, format='PNG')
                buf.seek(0)
                return send_file(buf, mimetype='image/png')
        except Exception as e:
            return f"Preview error: {e}", 500
    return "Not supported", 400


# ============ SAVE EDITS ============
@app.route('/save', methods=['POST'])
def save_edits():
    data = request.json
    file_id = data['file_id']
    ext = data['ext']
    edits = data.get('edits', [])
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.{ext}")
    output_path = os.path.join(app.config['EDITED_FOLDER'], f"{file_id}_edited.{ext}")

    if ext == 'pdf':
        edits_by_page = {}
        for edit in edits:
            edits_by_page.setdefault(edit['page'], []).append(edit)

        reader = PdfReader(filepath)
        writer = PdfWriter()
        for page_num, page in enumerate(reader.pages):
            if page_num in edits_by_page:
                w = float(page.mediabox.width)
                h = float(page.mediabox.height)
                packet = BytesIO()
                c = canvas.Canvas(packet, pagesize=(w, h))
                for edit in edits_by_page[page_num]:
                    c.setFillColorRGB(1, 1, 1)
                    c.rect(edit['x'], h - edit['y2'],
                           edit['x2'] - edit['x'], edit['y2'] - edit['y'],
                           fill=1, stroke=0)
                    c.setFillColorRGB(0, 0, 0)
                    c.setFont("Helvetica", edit.get('font_size', 11))
                    c.drawString(edit['x'], h - edit['y2'] + 2, edit['new_text'])
                c.save()
                packet.seek(0)
                overlay = PdfReader(packet)
                page.merge_page(overlay.pages[0])
            writer.add_page(page)
        with open(output_path, 'wb') as f:
            writer.write(f)

    elif ext in ('docx', 'doc'):
        doc = Document(filepath)
        for edit in edits:
            idx = edit['index']
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                if para.runs:
                    para.runs[0].text = edit['new_text']
                    for run in para.runs[1:]:
                        run.text = ''
                else:
                    para.text = edit['new_text']
        doc.save(output_path)

    return jsonify({'success': True, 'download_url': f'/download/{file_id}/{ext}'})


# ============ RE-ANALYZE ============
@app.route('/reanalyze/<file_id>/<ext>', methods=['POST'])
def reanalyze(file_id, ext):
    edited_path = os.path.join(app.config['EDITED_FOLDER'], f"{file_id}_edited.{ext}")
    original_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.{ext}")
    filepath = edited_path if os.path.exists(edited_path) else original_path

    analysis_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{file_id}.json")
    if not os.path.exists(analysis_path):
        return jsonify({'error': 'No previous analysis found'}), 404

    with open(analysis_path, 'r', encoding='utf-8') as f:
        prev = json.load(f)

    resume_text = extract_text_from_resume(filepath, ext)
    new_analysis = analyze_resume(resume_text, prev['role'], prev['job_description'])

    if 'error' in new_analysis:
        return jsonify(new_analysis), 500

    prev['analysis'] = new_analysis
    with open(analysis_path, 'w', encoding='utf-8') as f:
        json.dump(prev, f, indent=2)

    return jsonify({'success': True, 'analysis': new_analysis})


# ============ DOWNLOAD ============
@app.route('/download/<file_id>/<ext>')
def download_file(file_id, ext):
    output_path = os.path.join(app.config['EDITED_FOLDER'], f"{file_id}_edited.{ext}")
    if not os.path.exists(output_path):
        # If user hasn't edited yet, return original
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}.{ext}")
        if not os.path.exists(output_path):
            return "File not found", 404
    return send_file(output_path, as_attachment=True, download_name=f"updated_resume.{ext}")


if __name__ == '__main__':
    app.run(debug=True, port=5000)